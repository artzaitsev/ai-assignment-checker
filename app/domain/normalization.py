from __future__ import annotations

from io import BytesIO
import re
from pathlib import PurePosixPath

from app.domain.errors import NormalizationParseError, UnsupportedFormatError

SUPPORTED_FORMATS: tuple[str, ...] = (".txt", ".md", ".docx", ".pdf")


def extension_from_artifact_ref(*, artifact_ref: str) -> str:
    key = artifact_ref.removeprefix("stub://")
    suffix = PurePosixPath(key).suffix.lower()
    if suffix in SUPPORTED_FORMATS:
        return suffix
    raise UnsupportedFormatError(f"unsupported extension: {suffix or '<none>'}")


def parse_payload_to_text(*, extension: str, payload: bytes) -> str:
    try:
        if extension in (".txt", ".md"):
            return _parse_text(payload)
        if extension == ".docx":
            return _parse_docx(payload)
        if extension == ".pdf":
            return _parse_pdf(payload)
    except UnsupportedFormatError:
        raise
    except Exception as exc:
        raise NormalizationParseError(str(exc)) from exc

    raise UnsupportedFormatError(f"unsupported extension: {extension}")


def to_unified_markdown(*, text: str) -> str:
    normalized = text.replace("\x00", " ")
    normalized = re.sub(r"\r\n?", "\n", normalized)
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _parse_text(payload: bytes) -> str:
    try:
        return payload.decode("utf-8")
    except UnicodeDecodeError:
        return payload.decode("latin-1", errors="replace")


def _parse_docx(payload: bytes) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:  
        raise NormalizationParseError("python-docx is not installed") from exc

    doc = Document(BytesIO(payload))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        value = paragraph.text.strip()
        if value:
            parts.append(value)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _parse_pdf(payload: bytes) -> str:
    try:
        import pdfplumber
    except ModuleNotFoundError as exc:
        raise NormalizationParseError("pdfplumber is not installed") from exc

    pages: list[str] = []
    with pdfplumber.open(BytesIO(payload)) as pdf:
        for page in pdf.pages:
            value = (page.extract_text() or "").strip()
            if value:
                pages.append(value)
    return "\n\n".join(pages)
