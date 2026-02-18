from __future__ import annotations

import os
import re
import shutil
import zipfile
from pathlib import Path
from typing import Tuple, Union, List, Dict, Optional

import pandas as pd
from docx import Document

# PDF
try:
    import pdfplumber
    _HAS_PDFPLUMBER = True
except Exception:
    _HAS_PDFPLUMBER = False

# OCR
try:
    import pytesseract
    from PIL import Image
    _HAS_OCR = True
except Exception:
    _HAS_OCR = False


_TEXT_EXT = {".txt", ".md", ".sql"}
_DOCX_EXT = {".docx"}
_PDF_EXT = {".pdf"}
_IMG_EXT = {".png", ".jpg", ".jpeg", ".webp"}
_XLS_EXT = {".xlsx", ".xls"}


def _clean_text(s: str) -> str:
    s = s.replace("\x00", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return path.read_text(encoding="latin-1", errors="replace")


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if (c.text or "").strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    if not _HAS_PDFPLUMBER:
        raise RuntimeError("pdfplumber not installed. pip install pdfplumber")
    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = (page.extract_text() or "").strip()
            if t:
                parts.append(t)
    return "\n\n".join(parts)
'''
def _read_pdf_with_fallback_ocr(path: Path, ocr_lang: str = "rus+eng") -> str:
    if not _HAS_PDFPLUMBER:
        raise RuntimeError("pdfplumber не установлен. pip install pdfplumber")

    parts = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            try:
                t = (page.extract_text() or "").strip()
                if t:
                    parts.append(t)
                    continue
            # упали на шрифтах
            except Exception:
                pass  

            # OCR fallback
            if not _HAS_OCR:
                parts.append(f"[PDF_PAGE_OCR_SKIPPED page={i}: OCR not available]")
                continue

            img = page.to_image(resolution=200).original  # PIL image
            ocr_text = pytesseract.image_to_string(img, lang=ocr_lang).strip()
            if ocr_text:
                parts.append(ocr_text)
            else:
                parts.append(f"[PDF_PAGE_EMPTY page={i}]")

    return "\n\n".join(parts)
'''


def _read_excel(path: Path) -> str:
    xls = pd.read_excel(path, sheet_name=None, header=None)
    parts = []
    for sheet, df in xls.items():
        parts.append(f"=== SHEET: {sheet} ===")
        dff = df.fillna("")
        for _, row in dff.iterrows():
            line = " | ".join(str(x).strip() for x in row.tolist() if str(x).strip())
            if line:
                parts.append(line)
    return "\n".join(parts)


def _resolve_tesseract_cmd(tesseract_cmd: Optional[Union[str, Path]] = None) -> Optional[str]:
    if tesseract_cmd:
        p = Path(tesseract_cmd)
        return str(p) if p.exists() else str(tesseract_cmd)  
    env = os.getenv("TESSERACT_CMD")
    if env:
        return env

    found = shutil.which("tesseract")
    return found  

def _read_image_ocr(path: Path, lang: str, tesseract_cmd: Optional[str]) -> str:
    if not _HAS_OCR:
        raise RuntimeError("OCR unavailable. Install: pip install pytesseract pillow")

    if not tesseract_cmd:
        raise RuntimeError(
            "tesseract is not installed or not found. "
            "Pass tesseract_cmd=... or set env TESSERACT_CMD or add tesseract to PATH."
        )

    import pytesseract  
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    img = Image.open(path)
    return pytesseract.image_to_string(img, lang=lang)


def _read_any(path: Path, ocr_lang: str, tesseract_cmd: Optional[str]) -> Tuple[str, str]:
    ext = path.suffix.lower()
    if ext in _TEXT_EXT:
        return _read_text_file(path), "text"
    if ext in _DOCX_EXT:
        return _read_docx(path), "docx"
    if ext in _PDF_EXT:
        return _read_pdf(path), "pdf"
    if ext in _XLS_EXT:
        return _read_excel(path), "excel"
    if ext in _IMG_EXT:
        return _read_image_ocr(path, lang=ocr_lang, tesseract_cmd=tesseract_cmd), "ocr"
    raise RuntimeError(f"Unsupported extension: {ext}")


def zip_answers_to_df(
    zip_path: Union[str, Path],
    *,
    ocr_lang: str = "rus+eng",
    tesseract_cmd: Optional[Union[str, Path]] = None,
) -> pd.DataFrame:
    
    zip_path = Path(zip_path)
    assert zip_path.exists(), f"Zip not found: {zip_path}"

    out_dir = zip_path.parent / f"_unzipped_{zip_path.stem}"
    out_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(zip_path, "r") as z:
        z.extractall(out_dir)

    tess_cmd = _resolve_tesseract_cmd(tesseract_cmd)

    files = [p for p in out_dir.rglob("*") if p.is_file()]

    rows: List[Dict] = []
    for i, f in enumerate(sorted(files, key=lambda p: p.name.lower()), start=1):
        try:
            text, parser = _read_any(f, ocr_lang=ocr_lang, tesseract_cmd=tess_cmd)
            text = _clean_text(text)
            err = None
        except Exception as e:
            text, parser = "", "error"
            err = str(e)

        rows.append({
            "candidate_id": i,
            "file_name": f.name,
            #"file_relpath": str(f.relative_to(out_dir)),
            "ext": f.suffix.lower(),
            "parser": parser,
            "text": text,
            "error": err,
        })

    return pd.DataFrame(rows)
